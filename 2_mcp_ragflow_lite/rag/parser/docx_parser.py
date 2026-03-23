"""
DOCX 解析器
"""
import logging
import re
from typing import Optional

from rag.parser.base import BaseParser
from common.registry import parser_registry

logger = logging.getLogger(__name__)


@parser_registry.register(".docx")
@parser_registry.register(".doc")
class DocxParser(BaseParser):
    """DOCX/DOC 文档解析器"""

    name = "DOCX Parser"
    extensions = [".docx", ".doc"]

    def parse(self, filename: str, binary: Optional[bytes] = None):
        """解析 DOCX 文件，返回 (sections, tables)"""
        from docx import Document
        import io

        try:
            if binary:
                doc = Document(io.BytesIO(binary))
            else:
                doc = Document(filename)
        except Exception as e:
            logger.error(f"Failed to open DOCX {filename}: {e}")
            return [], []

        sections = []
        tables = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            tag = "text"
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if "heading" in style_name or "title" in style_name:
                    tag = "title"
            sections.append((text, tag))

        # 提取表格
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                header = rows[0]
                html = "<table>\n<tr>" + "".join(f"<th>{h}</th>" for h in header) + "</tr>\n"
                for row in rows[1:]:
                    html += "<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>\n"
                html += "</table>"
                tables.append(html)

        return sections, tables


# 向后兼容: 模块级 parse() 函数
_instance = DocxParser()
parse = _instance.parse
